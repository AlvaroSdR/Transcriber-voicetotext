import speech_recognition as sr
import docx
from pydub import AudioSegment
import os

def transcribe_audio(audio_file, output_file):
    """
    Transcribe the given audio file and append the recognized text to the output file.
    """
    r = sr.Recognizer()
    with sr.AudioFile(audio_file) as source:
        audio = r.record(source)
        text = r.recognize_google(audio)
        # Append the recognized text to the output file
        doc = docx.Document()
        if os.path.exists(output_file):
            doc = docx.Document(output_file)
        doc.add_paragraph(text)
        doc.add_paragraph("")  # Add a blank paragraph
        doc.save(output_file)

try:
    # Load the MP3 audio file
    audio_file = AudioSegment.from_file("notes.mp3", format="mp3")

    # Check the size of the MP3 file
    file_size = os.path.getsize("notes.mp3") / (1024 * 1024)  # File size in MB
    if file_size > 4:  # If file size exceeds 4 MB, split into smaller chunks
        chunk_size = 2  # Split into 2 MB chunks
        total_duration = audio_file.duration_seconds  # Total duration of the audio file in seconds
        num_chunks = int(total_duration // (chunk_size * 60)) + 1  # Number of chunks needed
        for i in range(num_chunks):
            chunk_start = i * chunk_size * 60 * 1000  # Start time of the chunk in milliseconds
            chunk_end = min((i + 1) * chunk_size * 60 * 1000, audio_file.duration_seconds * 1000)  # End time of the chunk in milliseconds
            chunk = audio_file[chunk_start:chunk_end]  # Extract the chunk
            chunk.export(f"notes_{i}.wav", format="wav")  # Export each chunk as a separate WAV file
            transcribe_audio(f"notes_{i}.wav", "text.docx")  # Transcribe each chunk and append to the output file
    else:
        # Convert the audio file to WAV format
        audio_file.export("notes.wav", format="wav")

        # Transcribe the audio file and save to output file
        transcribe_audio("notes.wav", "text.docx")

    print("Text has been saved to text.docx")
except sr.UnknownValueError:
    print("Speech recognition could not understand audio")
except sr.RequestError:
    print("Speech recognition request failed")
except Exception as e:
    print(f"Error: {e}")
